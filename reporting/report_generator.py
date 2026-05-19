from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPORTS_PATH = Path("data/reports")


def add_title(document, text):

    title = document.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(22)


def add_section(document, title, text):

    document.add_heading(title, level=2)

    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_metric_table(document, metrics):

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Métrica"
    header[1].text = "Resultado"

    for key, value in metrics.items():

        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)


def add_image(document, image_path, caption):

    image_path = Path(image_path)

    if image_path.exists():

        document.add_picture(
            str(image_path),
            width=Inches(6.2)
        )

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_p = document.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = caption_p.add_run(caption)
        run.italic = True


def generate_report(
    df,
    network_summary,
    channel_summary,
    bssid_summary,
    anomalies_df
):

    print("Gerando relatório DOCX profissional...")

    document = Document()

    # =========================
    # MÉTRICAS PRINCIPAIS
    # =========================

    total_records = len(df)
    unique_ssids = df["ssid"].nunique()
    unique_bssids = df["bssid"].nunique()

    avg_rssi = round(df["rssi"].mean(), 2)
    best_rssi = int(df["rssi"].max())
    worst_rssi = int(df["rssi"].min())

    avg_quality = round(
        df["signal_quality_score"].mean(),
        2
    )

    top_network = (
        network_summary
        .sort_values("detections", ascending=False)
        .iloc[0]
    )

    strongest_network = (
        network_summary
        .sort_values("avg_rssi", ascending=False)
        .iloc[0]
    )

    most_unstable_network = (
        network_summary
        .sort_values("rssi_std", ascending=False)
        .iloc[0]
    )

    most_used_channel = (
        channel_summary
        .sort_values("detections", ascending=False)
        .iloc[0]
    )

    dominant_bssid = (
        bssid_summary
        .sort_values("presence_score", ascending=False)
        .iloc[0]
    )

    most_stable_bssid = (
        bssid_summary
        .sort_values("stability_score", ascending=False)
        .iloc[0]
    )

    anomaly_count = len(anomalies_df)

    # =========================
    # CAPA
    # =========================

    add_title(
        document,
        "WiFi Scope - Relatório de Inteligência Wireless"
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "Análise RF baseada em ESP8266, Google Sheets e Python"
    )

    subtitle_run.italic = True

    date_text = datetime.now().strftime(
        "%d/%m/%Y %H:%M:%S"
    )

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.add_run(f"Gerado em: {date_text}")

    document.add_page_break()

    # =========================
    # RESUMO EXECUTIVO
    # =========================

    executive_summary = f"""
Durante o período analisado, o sistema WiFi Scope registrou {total_records} detecções de redes Wi-Fi, identificando {unique_ssids} SSIDs distintos e {unique_bssids} BSSIDs/APs únicos.

A intensidade média geral do sinal foi de {avg_rssi} dBm, com melhor sinal registrado em {best_rssi} dBm e pior sinal registrado em {worst_rssi} dBm.

A rede mais presente no ambiente foi "{top_network['ssid']}", com {top_network['detections']} detecções. A rede com melhor sinal médio foi "{strongest_network['ssid']}", apresentando RSSI médio de {strongest_network['avg_rssi']} dBm.

O canal mais utilizado foi o canal {int(most_used_channel['channel'])}, indicando maior concentração de atividade RF nessa faixa.

Foram detectadas {anomaly_count} ocorrências classificadas como anomalias, considerando sinais muito fracos, sinais extremamente fortes e redes detectadas apenas uma vez.
"""

    add_section(
        document,
        "1. Resumo Executivo",
        executive_summary
    )

    metrics = {
        "Total de registros": total_records,
        "SSIDs únicos": unique_ssids,
        "BSSIDs/APs únicos": unique_bssids,
        "RSSI médio": f"{avg_rssi} dBm",
        "Melhor RSSI": f"{best_rssi} dBm",
        "Pior RSSI": f"{worst_rssi} dBm",
        "Qualidade média": f"{avg_quality}/100",
        "Canal mais usado": int(most_used_channel["channel"]),
        "Anomalias detectadas": anomaly_count
    }

    add_metric_table(document, metrics)

    # =========================
    # GRÁFICO 1
    # =========================

    add_section(
        document,
        "2. Redes Mais Detectadas",
        f"""
Este gráfico mostra quais redes apareceram com maior frequência durante o monitoramento.

A rede mais recorrente foi "{top_network['ssid']}", associada ao BSSID {top_network['bssid']}, com {top_network['detections']} detecções.

Alta recorrência indica presença constante no ambiente RF monitorado, o que pode representar um roteador próximo, estável ou dominante na região.
"""
    )

    add_image(
        document,
        "data/reports/top_networks.png",
        "Figura 1 - Top redes mais detectadas"
    )

    # =========================
    # GRÁFICO 2
    # =========================

    add_section(
        document,
        "3. Redes com Melhor Sinal Médio",
        f"""
A análise de sinal médio permite identificar quais redes apresentaram maior intensidade de recepção.

A rede com melhor sinal médio foi "{strongest_network['ssid']}", com RSSI médio de {strongest_network['avg_rssi']} dBm.

Em redes Wi-Fi, valores menos negativos representam sinais melhores. Por exemplo, -30 dBm indica sinal muito forte, enquanto valores próximos de -85 dBm indicam sinal fraco.
"""
    )

    add_image(
        document,
        "data/reports/strongest_networks.png",
        "Figura 2 - Redes com melhor sinal médio"
    )

    # =========================
    # GRÁFICO 3
    # =========================

    add_section(
        document,
        "4. Redes Mais Instáveis",
        f"""
Este gráfico utiliza o desvio padrão do RSSI para identificar redes com maior variação de sinal.

A rede mais instável foi "{most_unstable_network['ssid']}", com desvio padrão de RSSI igual a {most_unstable_network['rssi_std']}.

Variações elevadas podem indicar interferência, distância elevada, obstáculos físicos, mudanças de propagação ou baixa estabilidade do ponto de acesso.
"""
    )

    add_image(
        document,
        "data/reports/most_unstable_networks.png",
        "Figura 3 - Redes mais instáveis"
    )

    # =========================
    # GRÁFICO 4
    # =========================

    add_section(
        document,
        "5. Distribuição de Redes por Canal",
        f"""
A distribuição por canal permite entender como o espectro de 2.4 GHz está sendo ocupado.

O canal mais utilizado foi o canal {int(most_used_channel['channel'])}, com {int(most_used_channel['detections'])} detecções.

Concentração excessiva em um canal pode aumentar interferência, colisões e retransmissões, afetando a qualidade geral da comunicação wireless.
"""
    )

    add_image(
        document,
        "data/reports/channel_distribution.png",
        "Figura 4 - Distribuição de redes por canal"
    )

    # =========================
    # GRÁFICO 5
    # =========================

    add_section(
        document,
        "6. RSSI ao Longo do Tempo",
        """
Este gráfico mostra a evolução temporal do RSSI das redes mais presentes.

Ele permite observar variações de intensidade, estabilidade, quedas de sinal e comportamento temporal das principais redes detectadas.

Esse tipo de análise é importante para identificar padrões de propagação e possíveis momentos de degradação do ambiente RF.
"""
    )

    add_image(
        document,
        "data/reports/rssi_timeline.png",
        "Figura 5 - RSSI ao longo do tempo"
    )

    # =========================
    # GRÁFICO 6
    # =========================

    add_section(
        document,
        "7. Score de Carga por Canal",
        """
O Channel Load Score combina quantidade de detecções e número de APs únicos por canal.

Esse indicador ajuda a identificar canais com maior concentração de atividade, funcionando como uma métrica simples de ocupação espectral.

Quanto maior o score, maior tende a ser a carga RF naquele canal.
"""
    )

    add_image(
        document,
        "data/reports/channel_load_score.png",
        "Figura 6 - Score de carga por canal"
    )

    # =========================
    # GRÁFICO 7
    # =========================

    add_section(
        document,
        "8. RSSI Médio por Canal",
        """
Este gráfico mostra a intensidade média dos sinais detectados em cada canal.

Ele ajuda a entender se determinados canais concentram sinais mais fortes ou mais fracos.

Canais com muitos sinais fortes próximos podem representar maior competição pelo meio wireless.
"""
    )

    add_image(
        document,
        "data/reports/avg_rssi_by_channel.png",
        "Figura 7 - RSSI médio por canal"
    )

    # =========================
    # GRÁFICO 8
    # =========================

    add_section(
        document,
        "9. Presença por BSSID/AP",
        f"""
A análise por BSSID é mais precisa que a análise por SSID, pois o BSSID identifica fisicamente um ponto de acesso específico.

O AP mais presente foi o BSSID {dominant_bssid['bssid']}, associado ao SSID "{dominant_bssid['ssid']}", com presence score de {dominant_bssid['presence_score']}%.

Esse indicador ajuda a identificar quais equipamentos realmente dominam o ambiente RF.
"""
    )

    add_image(
        document,
        "data/reports/bssid_presence.png",
        "Figura 8 - Presença por BSSID/AP"
    )

    # =========================
    # GRÁFICO 9
    # =========================

    add_section(
        document,
        "10. Estabilidade por BSSID/AP",
        f"""
A estabilidade por BSSID avalia o comportamento individual de cada ponto de acesso.

O BSSID mais estável foi {most_stable_bssid['bssid']}, associado ao SSID "{most_stable_bssid['ssid']}", com stability score de {most_stable_bssid['stability_score']}%.

Esse tipo de análise é essencial para diferenciar redes apenas fortes de redes realmente consistentes ao longo do tempo.
"""
    )

    add_image(
        document,
        "data/reports/bssid_stability.png",
        "Figura 9 - Estabilidade por BSSID/AP"
    )

    # =========================
    # GRÁFICO 10
    # =========================

    add_section(
        document,
        "11. Tipos de Anomalias Detectadas",
        f"""
O sistema identificou {anomaly_count} eventos classificados como anomalias.

As anomalias podem indicar redes detectadas apenas uma vez, sinais extremamente fortes ou sinais muito fracos.

Essa análise é útil para identificar eventos raros, redes temporárias, alterações no ambiente ou sinais fora do padrão esperado.
"""
    )

    add_image(
        document,
        "data/reports/anomaly_types.png",
        "Figura 10 - Distribuição dos tipos de anomalias"
    )

    # =========================
    # GRÁFICO 11
    # =========================

    add_section(
        document,
        "12. Distribuição da Qualidade de Sinal",
        """
A distribuição de qualidade classifica cada detecção com base no RSSI convertido em uma escala positiva de 0 a 100.

Essa abordagem facilita a interpretação dos dados, já que RSSI em dBm é negativo e pode ser contraintuitivo.

A classificação permite observar rapidamente se o ambiente é majoritariamente excelente, bom, moderado, ruim ou crítico.
"""
    )

    add_image(
        document,
        "data/reports/quality_distribution.png",
        "Figura 11 - Distribuição da qualidade de sinal"
    )

    # =========================
    # CONCLUSÃO DINÂMICA
    # =========================

    if avg_quality >= 85:
        quality_text = "O ambiente apresentou qualidade geral excelente."
    elif avg_quality >= 70:
        quality_text = "O ambiente apresentou boa qualidade geral de sinal."
    elif avg_quality >= 50:
        quality_text = "O ambiente apresentou qualidade moderada, com pontos de atenção."
    else:
        quality_text = "O ambiente apresentou qualidade baixa, exigindo investigação."

    if anomaly_count > total_records * 0.25:
        anomaly_text = "A proporção de anomalias foi elevada, indicando variações importantes no ambiente RF."
    else:
        anomaly_text = "A proporção de anomalias foi controlada em relação ao total de registros."

    conclusion = f"""
Com base nos dados coletados, o WiFi Scope conseguiu mapear o comportamento wireless do ambiente com boa granularidade, analisando redes, BSSIDs, canais, intensidade de sinal, estabilidade e anomalias.

{quality_text}

{anomaly_text}

O projeto demonstra uma arquitetura funcional de observabilidade wireless, combinando ESP8266, Google Sheets e Python para transformar dados RF em informação técnica útil.

A próxima evolução natural seria adicionar dashboard interativo, alertas automáticos e análise preditiva.
"""

    add_section(
        document,
        "13. Conclusão Técnica",
        conclusion
    )

    # =========================
    # SALVAR DOCX
    # =========================

    today = datetime.now().strftime("%d-%m-%Y")
    timestamp = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")

    folder = Path(
    f"C:/Users/rferr/OneDrive/Documentos/wifi scope docs/{today}"
    )


    folder.mkdir(
        parents=True,
        exist_ok=True
    )

    output_path = (
        folder / f"WiFi_Scope_Report_{timestamp}.docx"
    )

    document.save(output_path)

    print("Relatório DOCX gerado:")
    print(output_path)

    return output_path